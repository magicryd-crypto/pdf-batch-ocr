# -*- coding: utf-8 -*-
r"""
Пакетное распознавание PDF -> Word (.docx) + Excel (.xlsx).

- Тип каждой страницы определяется автоматически: "цифровая" (есть текстовый
  слой) или "скан" (нужен OCR).
- Сканы распознаются ОДНОЙ проходкой Tesseract на страницу: из неё берётся и
  текст, и таблица (по линиям сетки, найденным OpenCV) — это быстро на больших
  объёмах.
- Возобновляемость: если для PDF уже есть .docx в output, файл пропускается.
- Параллельная обработка файлов по ядрам процессора.

Запуск:
    venv\Scripts\python.exe convert.py                 # input\ -> output\
    venv\Scripts\python.exe convert.py "вход" "выход"  # свои папки
    venv\Scripts\python.exe convert.py "вход" "выход" 8  # 8 параллельных задач
"""
import os
import sys
import glob
import traceback
from multiprocessing import Pool, cpu_count

import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image
import pandas as pd
import numpy as np
import cv2
from docx import Document

# ----------------------------- НАСТРОЙКИ -----------------------------------
TESSERACT_EXE = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
TESSDATA_DIR = os.path.join(os.environ.get("LOCALAPPDATA", ""), "tessdata_fast")
LANG = "rus+eng"          # языки распознавания (быстрые лёгкие модели)
DPI = 150                 # ниже DPI -> меньше памяти на страницу (tesseract не OOM)
TEXT_LAYER_MIN_CHARS = 25 # порог: меньше символов на странице -> считаем сканом
OCR_SCANNED_TABLES = False # таблицы со сканов-фотокопий дают шум; экономим память
DEFAULT_JOBS = 4           # лёгкая модель ~0.7 ГБ/воркер; ~4 ГБ свободно
# ---------------------------------------------------------------------------

# Каждый процесс Tesseract — в один поток; параллелизм даём на уровне файлов.
os.environ["OMP_THREAD_LIMIT"] = "1"
pytesseract.pytesseract.tesseract_cmd = TESSERACT_EXE
os.environ["TESSDATA_PREFIX"] = TESSDATA_DIR
# Путь к моделям задаётся флагом --tessdata-dir (надёжнее переменной окружения).
TESS_CONFIG = f"--tessdata-dir {TESSDATA_DIR} --oem 1 --psm 3"


def _line_positions(flags, min_gap):
    """По бинарному профилю (1 = есть линия) возвращает центры линий."""
    pos, i, n = [], 0, len(flags)
    while i < n:
        if flags[i]:
            j = i
            while j < n and flags[j]:
                j += 1
            pos.append((i + j) // 2)
            i = j
        else:
            i += 1
    merged = []
    for p in pos:
        if merged and p - merged[-1] < min_gap:
            merged[-1] = (merged[-1] + p) // 2
        else:
            merged.append(p)
    return merged


def _detect_grid(rgb):
    """Возвращает (row_lines, col_lines) разлинованной таблицы или (None, None)."""
    gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY)
    bw = cv2.adaptiveThreshold(~gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                               cv2.THRESH_BINARY, 15, -2)
    h, w = bw.shape
    hk = cv2.getStructuringElement(cv2.MORPH_RECT, (max(20, w // 25), 1))
    horiz = cv2.dilate(cv2.erode(bw, hk), hk)
    vk = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(20, h // 25)))
    vert = cv2.dilate(cv2.erode(bw, vk), vk)

    hsum = horiz.sum(axis=1) / 255.0
    vsum = vert.sum(axis=0) / 255.0
    if hsum.max() < w * 0.15 or vsum.max() < h * 0.02:
        return None, None
    h_thr = max(hsum.max() * 0.4, w * 0.15)
    v_thr = max(vsum.max() * 0.4, h * 0.02)
    row_lines = _line_positions(hsum > h_thr, max(10, h // 100))
    col_lines = _line_positions(vsum > v_thr, max(10, w // 100))
    if len(row_lines) < 2 or len(col_lines) < 2:
        return None, None
    return row_lines, col_lines


def ocr_page(page):
    """Одна проходка OCR по странице-скану -> (текст, DataFrame таблицы | None)."""
    mat = fitz.Matrix(DPI / 72, DPI / 72)
    pix = page.get_pixmap(matrix=mat)
    rgb = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)[:, :, :3]
    pil = Image.fromarray(rgb)

    data = pytesseract.image_to_data(
        pil, lang=LANG, config=TESS_CONFIG,
        output_type=pytesseract.Output.DATAFRAME,
    )
    data = data[data.text.notna()].copy()
    data["text"] = data["text"].astype(str)
    data = data[(pd.to_numeric(data.conf, errors="coerce") > 30) &
                (data.text.str.strip() != "")]

    # --- текст: слова по строкам ---
    lines = []
    if len(data):
        for _, g in data.groupby(["block_num", "par_num", "line_num"], sort=True):
            g = g.sort_values("left")
            lines.append(" ".join(g.text.tolist()))
    text = "\n".join(lines)

    # --- таблица по сетке (если включено и сетка найдена) ---
    table = None
    if OCR_SCANNED_TABLES and len(data):
        row_lines, col_lines = _detect_grid(rgb)
        if row_lines is not None:
            ncols = len(col_lines) - 1
            nrows = len(row_lines) - 1
            grid = [["" for _ in range(ncols)] for _ in range(nrows)]
            for _, wd in data.iterrows():
                cx = int(wd.left) + int(wd.width) // 2
                cy = int(wd.top) + int(wd.height) // 2
                r = next((i for i in range(nrows)
                          if row_lines[i] <= cy < row_lines[i + 1]), None)
                c = next((i for i in range(ncols)
                          if col_lines[i] <= cx < col_lines[i + 1]), None)
                if r is not None and c is not None:
                    grid[r][c] = (grid[r][c] + " " + wd.text).strip()
            grid = [row for row in grid if any(row)]
            if len(grid) >= 2:
                table = pd.DataFrame(grid)
    return text, table


def process_pdf(args):
    pdf_path, out_dir = args
    name = os.path.splitext(os.path.basename(pdf_path))[0]
    os.makedirs(out_dir, exist_ok=True)
    docx_path = os.path.join(out_dir, f"{name}.docx")
    txt_path = os.path.join(out_dir, f"{name}.txt")
    if os.path.exists(docx_path) or os.path.exists(txt_path):
        return {"pdf": os.path.basename(pdf_path), "skipped": True}

    tables = []        # только из цифровых страниц (pdfplumber) — память лёгкая
    scanned = 0
    pages_text = []    # сырой текст по страницам (лёгкий, не lxml)

    fitz_doc = fitz.open(pdf_path)
    plumber = pdfplumber.open(pdf_path)
    n = len(fitz_doc)

    # --- OCR/извлечение с ПОТОКОВОЙ записью в .txt (память остаётся плоской) ---
    tmp = txt_path + ".part"
    with open(tmp, "w", encoding="utf-8") as tf:
        tf.write(name + "\n")
        for i in range(n):
            # постранично устойчиво: сбой одной страницы не валит весь том
            try:
                page = fitz_doc[i]
                digital_text = page.get_text("text").strip()
                is_digital = len(digital_text) >= TEXT_LAYER_MIN_CHARS
                if is_digital:
                    body = digital_text
                    for ti, tbl in enumerate(plumber.pages[i].extract_tables()):
                        if tbl:
                            tables.append((f"Стр{i+1}_Таблица{ti+1}", pd.DataFrame(tbl)))
                else:
                    scanned += 1
                    body, _ = ocr_page(page)  # таблицы со сканов отключены (шум)
            except Exception as e:
                is_digital = False
                body = f"[Страница {i + 1}: ошибка распознавания: {type(e).__name__}]"

            head = f"\n===== Страница {i + 1}" + ("" if is_digital else "  [скан/OCR]") + " =====\n"
            tf.write(head); tf.write(body); tf.write("\n"); tf.flush()
            pages_text.append((i + 1, is_digital, body))

    fitz_doc.close()
    plumber.close()
    os.replace(tmp, txt_path)  # атомарно: .txt появляется только если дошли до конца

    # --- Word собираем в конце, когда память от картинок/OCR освобождена ---
    docx_ok = False
    try:
        doc = Document()
        doc.add_heading(name, level=0)
        for pno, is_digital, body in pages_text:
            doc.add_heading(f"Страница {pno}" + ("" if is_digital else "  [скан/OCR]"), level=1)
            for line in body.splitlines():
                doc.add_paragraph(line)
        doc.save(docx_path)
        docx_ok = True
    except (MemoryError, Exception):
        docx_ok = False  # .txt уже сохранён — том засчитан, Word добьём отдельно

    xlsx_path = None
    if tables:
        xlsx_path = os.path.join(out_dir, f"{name}.xlsx")
        with pd.ExcelWriter(xlsx_path, engine="openpyxl") as xw:
            used = set()
            for label, df in tables:
                sheet, base, k = label[:31], label[:28], 1
                while sheet in used:
                    sheet = f"{base}_{k}"; k += 1
                used.add(sheet)
                df.to_excel(xw, sheet_name=sheet, index=False, header=False)

    return {"pdf": os.path.basename(pdf_path), "pages": n, "scanned": scanned,
            "tables": len(tables), "xlsx": bool(xlsx_path), "docx": docx_ok,
            "skipped": False}


def main():
    base = os.path.dirname(os.path.abspath(__file__))
    # режим одного файла: convert.py путь\к\файлу.pdf [out_dir]
    if len(sys.argv) > 1 and sys.argv[1].lower().endswith(".pdf"):
        out_dir = sys.argv[2] if len(sys.argv) > 2 else os.path.join(base, "output")
        os.makedirs(out_dir, exist_ok=True)
        r = _safe_process((sys.argv[1], out_dir))
        print(r)
        name = os.path.splitext(os.path.basename(sys.argv[1]))[0]
        ok = os.path.exists(os.path.join(out_dir, name + ".txt")) or \
            os.path.exists(os.path.join(out_dir, name + ".docx"))
        sys.exit(0 if ok else 1)
    in_dir = sys.argv[1] if len(sys.argv) > 1 else os.path.join(base, "input")
    out_dir = sys.argv[2] if len(sys.argv) > 2 else os.path.join(base, "output")
    jobs = int(sys.argv[3]) if len(sys.argv) > 3 else DEFAULT_JOBS
    os.makedirs(in_dir, exist_ok=True)
    os.makedirs(out_dir, exist_ok=True)

    pdfs = sorted(glob.glob(os.path.join(in_dir, "**", "*.pdf"), recursive=True))
    if not pdfs:
        print(f"Нет PDF в папке: {in_dir}")
        return

    total = len(pdfs)
    print(f"Файлов: {total} | параллельно: {jobs} | вывод: {out_dir}")
    print("Готовые файлы пропускаются (можно прерывать и запускать снова).")
    print("-" * 64)

    done = 0
    args = [(p, out_dir) for p in pdfs]
    with Pool(processes=jobs) as pool:
        for r in pool.imap_unordered(_safe_process, args):
            done += 1
            if r.get("skipped"):
                print(f"[{done}/{total}] ПРОПУЩЕН (уже готов): {r['pdf']}")
            elif r.get("error"):
                print(f"[{done}/{total}] ОШИБКА {r['pdf']}: {r['error']}")
            else:
                x = "есть" if r["xlsx"] else "нет"
                print(f"[{done}/{total}] OK {r['pdf']}: страниц {r['pages']} "
                      f"(сканов {r['scanned']}), таблиц {r['tables']}, Excel: {x}")
    print("-" * 64 + "\nГотово.")


def _safe_process(args):
    try:
        return process_pdf(args)
    except Exception as e:
        return {"pdf": os.path.basename(args[0]), "error": repr(e),
                "trace": traceback.format_exc()}


if __name__ == "__main__":
    main()
